"""
build_docx.py
Gabungkan semua file .md di folder laporan/ menjadi SATU dokumen Word (.docx)
dengan format rapi (cover, daftar isi, header/footer, code block styling, tabel).
"""

from pathlib import Path
import re
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import modul render diagram (non-aktif: flow diagram hanya di-render
# ke PNG di folder diagrams/, tidak di-embed ke DOCX)
try:
    from render_diagrams import (
        is_ascii_diagram, render_diagram_block,
        DIAGRAMS_DIR, estimate_docx_width
    )
    RENDER_DIAGRAMS = False   # ← DISABLED: diagram tetap ASCII di DOCX
except Exception as _e:
    print(f"[WARN] render_diagrams tidak tersedia: {_e}")
    RENDER_DIAGRAMS = False

# ─────────────────────────────────────────────────────────────────────────
# PATH
# ─────────────────────────────────────────────────────────────────────────
BASE  = Path(__file__).parent
OUT   = BASE / "Laporan_UAS_eTicketing_Helpdesk.docx"
OUT_FALLBACK = BASE / "Laporan_UAS_eTicketing_Helpdesk_new.docx"
FILES = [
    ("01-pendahuluan.md",   "BAB 1  PENDAHULUAN"),
    ("02-arsitektur.md",    "BAB 2  ARSITEKTUR SISTEM"),
    ("03-database.md",      "BAB 3  DATABASE"),
    ("04-api.md",           "BAB 4  BACKEND API"),
    ("05-uiux.md",          "BAB 5  UI / UX APLIKASI"),
    ("06-fitur.md",         "BAB 6  FITUR APLIKASI"),
    ("07-video-tutorial.md","BAB 7  VIDEO TUTORIAL"),
    ("08-panduan-deploy.md","BAB 8  PANDUAN DEPLOY & INSTALASI"),
    ("09-dokumentasi-kode.md","BAB 9  DOKUMENTASI KODE"),
]

# ─────────────────────────────────────────────────────────────────────────
# WARNA & STYLE
# ─────────────────────────────────────────────────────────────────────────
INK        = RGBColor(0x00, 0x00, 0x00)         # hitam pekat
YELLOW     = RGBColor(0xFF, 0xD7, 0x00)         # kuning Brutalism
PINK_NEON  = RGBColor(0xFF, 0x3D, 0x7F)
ORANGE     = RGBColor(0xFF, 0x7A, 0x00)
LIME       = RGBColor(0x73, 0xF7, 0x4D)
BLUE       = RGBColor(0x00, 0xC2, 0xFF)
KREM       = RGBColor(0xFF, 0xFB, 0xF0)
CODE_BG    = "F4F4F4"                            # abu terang untuk code block
SHADE      = "FFF6E0"                            # krem muda untuk alt row


# ─────────────────────────────────────────────────────────────────────────
# HELPER: shading cell
# ─────────────────────────────────────────────────────────────────────────
def _shade_cell(cell, fill_hex):
    tcPr  = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ─────────────────────────────────────────────────────────────────────────
# PARSER: super-sederhana (cukup untuk konten laporan UAS)
# Handles:
#   - #, ##, ### ...  → heading 1..6
#   - ``` code block (multi-line)
#   - | table | with separator line ─
#   - - / * bullet
#   - > blockquote
#   - **bold** *italic* `inline code`
#   - --- horizontal rule
# ─────────────────────────────────────────────────────────────────────────

class MdParser:
    def __init__(self, doc):
        self.doc = doc
        self.in_code = False
        self.code_lines = []
        self.code_lang  = ""
        self.code_hint  = ""   # nama section untuk naming PNG
        self.in_table = False
        self.table_rows = []
        # State: judul heading terakhir (untuk hint nama file diagram)
        self.last_heading = ""

    # ── util ─────────────────────────────────────────────────────────────
    def _apply_inline(self, run, text, is_code_block=False):
        """Tambahkan teks ke run, parsing bold/italic/inline-code sederhana."""
        # Tokenisasi bold **..**, italic *..*, code `..`
        pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
        pos = 0
        parts = pattern.split(text)
        for chunk in parts:
            if not chunk:
                continue
            if chunk.startswith('**') and chunk.endswith('**'):
                run.add_text(chunk[2:-2])
                run.bold = True
            elif chunk.startswith('*') and chunk.endswith('*') and len(chunk) > 2:
                run.add_text(chunk[1:-1])
                run.italic = True
            elif chunk.startswith('`') and chunk.endswith('`'):
                run.add_text(chunk[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                rPr = run._r.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'EFEFEF')
                rPr.append(shd)
            else:
                run.add_text(chunk)

    def _add_run(self, paragraph, text):
        run = paragraph.add_run()
        self._apply_inline(run, text)
        return run

    # ── block handlers ───────────────────────────────────────────────────
    def handle_heading(self, line):
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if not m:
            return False
        level = len(m.group(1))
        text  = m.group(2).strip()
        # Simpan sebagai hint untuk nama file diagram
        self.last_heading = re.sub(r'[^a-zA-Z0-9]+', '-', text.lower()).strip('-')
        # Map level: docx hanya punya heading 1..9 built-in
        style_name = f"Heading {min(level + 1, 4)}"  # geser: H1 → Heading 3, dst.
        # Override: bab pakai Heading 1, sub Heading 2, dst.
        level_map = {1: 1, 2: 2, 3: 3, 4: 4, 5: 5, 6: 6}
        style_name = f"Heading {level_map[level]}"
        p = self.doc.add_paragraph(style=style_name)
        p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.color.rgb = INK
        if level == 1:
            run.font.size = Pt(20)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(16)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(14)
            run.bold = True
        elif level == 4:
            run.font.size = Pt(12)
            run.bold = True
        else:
            run.font.size = Pt(11)
            run.bold = True
        return True

    def handle_code_fence(self, line):
        global _code_open
        m = re.match(r'^```(\w*)$', line.strip())
        if not m and not self.in_code:
            return False
        if not self.in_code:
            self.in_code = True
            self.code_lang = m.group(1) if m else ""
            self.code_hint = self.last_heading or "diagram"
            self.code_lines = []
            return True
        else:
            # Close code block
            self.in_code = False
            self._render_code_block('\n'.join(self.code_lines))
            self.code_lines = []
            return True

    def _render_code_block(self, code_text):
        # ── Deteksi: ASCII flow diagram? ───────────────────────────────
        # Aturan:
        #   • code_lang kosong ATAU tidak berisi keyword bahasa (dart/sql/js/py)
        #   • DAN text punya box-drawing chars / arrows
        langs_real = {"dart", "sql", "javascript", "js", "typescript", "ts",
                      "python", "py", "java", "kotlin", "kt", "yaml", "yml",
                      "json", "bash", "sh", "shell", "html", "css", "scss",
                      "xml", "ini", "env", "toml", "markdown", "md",
                      "txt", "text", "plaintext", "log"}
        is_real_code_lang = self.code_lang.lower() in langs_real

        if (RENDER_DIAGRAMS
                and not is_real_code_lang
                and is_ascii_diagram(code_text)):
            # Render ke PNG & embed sebagai image
            png_path = render_diagram_block(code_text, hint=self.code_hint)
            try:
                # Hitung width
                from PIL import Image as _Img
                with _Img.open(png_path) as _im:
                    w_px, h_px = _im.size
                max_w_in = estimate_docx_width(w_px)
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(10)
                run = p.add_run()
                run.add_picture(str(png_path), width=Inches(max_w_in))
                return
            except Exception as e:
                print(f"[WARN] Gagal embed diagram {png_path}: {e}")

        # ── Default: code block monospace dengan shading ──────────────
        # Paragraph tunggal dengan semua baris, font monospace, shading abu
        # Pakai tabel 1×1 untuk visual code block lebih rapi
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.rows[0].cells[0]
        _shade_cell(cell, CODE_BG)
        # Border tipis
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), '4')
            e.set(qn('w:color'), '888888')
            tcBorders.append(e)
        tcPr.append(tcBorders)

        # Hapus paragraph default
        cell.text = ''
        first = True
        for line in code_text.split('\n'):
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line if line else ' ')
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = INK
        doc_para = self.doc.add_paragraph()
        doc_para.paragraph_format.space_after = Pt(6)

    def handle_table(self, line):
        # Tabel markdown: | a | b |
        if not re.match(r'^\s*\|.+\|\s*$', line):
            return False
        # Deteksi apakah ini baris data atau baris pemisah (|---|---|)
        if re.match(r'^\s*\|[-:\s|]+\|\s*$', line):
            return True  # skip separator
        # Parse cells
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if not self.in_table:
            self.in_table = True
            self.table_rows = []
        self.table_rows.append(cells)
        return True

    def _render_table(self):
        if not self.table_rows:
            self.in_table = False
            return
        # Pastikan semua row sama lebar
        ncols = max(len(r) for r in self.table_rows)
        for r in self.table_rows:
            while len(r) < ncols:
                r.append('')

        tbl = self.doc.add_table(rows=len(self.table_rows), cols=ncols)
        tbl.style = 'Table Grid'
        tbl.autofit = True

        for ridx, row in enumerate(self.table_rows):
            for cidx, val in enumerate(row):
                cell = tbl.rows[ridx].cells[cidx]
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run()
                self._apply_inline(run, val)
                if ridx == 0:
                    # Header row: kuning Brutalism
                    run.bold = True
                    _shade_cell(cell, 'FFD700')
                else:
                    if ridx % 2 == 0:
                        _shade_cell(cell, SHADE)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Spacer
        self.doc.add_paragraph()

    def handle_image(self, alt: str, path: str):
        """Embed image dari path (relatif ke folder laporan/ atau absolut)."""
        # Resolve path
        if Path(path).is_absolute():
            full = Path(path)
        else:
            full = BASE / path
        if not full.exists():
            print(f"[WARN] Image tidak ditemukan: {full}")
            # Fallback: tampilkan alt text sebagai paragraph
            p = self.doc.add_paragraph()
            r = p.add_run(f"[Gambar: {alt or path}]")
            r.italic = True
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return
        try:
            from PIL import Image as _Img
            with _Img.open(full) as _im:
                w_px, h_px = _im.size
            max_w_in = estimate_docx_width(w_px)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run()
            run.add_picture(str(full), width=Inches(max_w_in))
            # Caption
            if alt:
                cp = self.doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(10)
                cr = cp.add_run(alt)
                cr.italic = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                self.doc.add_paragraph()
        except Exception as e:
            print(f"[WARN] Gagal embed image {full}: {e}")

    def handle_blockquote(self, line):
        m = re.match(r'^>\s?(.*)$', line)
        if not m:
            return False
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.right_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(4)
        # Italic + abu
        run = p.add_run('│ ' + m.group(1))
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        return True

    def handle_hr(self, line):
        if line.strip() == '---':
            _add_horizontal_rule(self.doc)
            return True
        return False

    def handle_bullet(self, line):
        m = re.match(r'^(\s*)([-*+])\s+(.+)$', line)
        if not m:
            return False
        indent_spaces = len(m.group(1))
        text = m.group(3)
        level = indent_spaces // 2
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
        p.paragraph_format.space_after = Pt(3)
        self._add_run(p, text)
        return True

    def handle_ordered(self, line):
        m = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if not m:
            return False
        indent_spaces = len(m.group(1))
        text = m.group(3)
        p = self.doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Cm(0.6 + indent_spaces * 0.3)
        p.paragraph_format.space_after = Pt(3)
        self._add_run(p, text)
        return True

    def handle_paragraph(self, line):
        # Skip kalau line kosong atau marker khusus
        if not line.strip():
            return False
        if self.handle_heading(line): return True
        if self.handle_blockquote(line): return True
        if self.handle_hr(line): return True
        if self.handle_bullet(line): return True
        if self.handle_ordered(line): return True
        if self.handle_table(line): return True
        if self.handle_code_fence(line): return True

        # Cek apakah line adalah bagian dari tabel aktif
        if self.in_table:
            return True  # skip, tunggu baris berikutnya

        # Cek apakah line adalah bagian dari code block aktif
        if self.in_code:
            self.code_lines.append(line)
            return True

        # Paragraph biasa
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        self._add_run(p, line)
        return True

    def feed(self, lines):
        i = 0
        n = len(lines)
        while i < n:
            line = lines[i]

            # Code block handling — multi-line state
            if self.in_code:
                if re.match(r'^```\s*$', line.strip()):
                    self.handle_code_fence('```')
                    i += 1
                    continue
                else:
                    self.code_lines.append(line)
                    i += 1
                    continue

            # Tabel — sedang dalam tabel tapi line sekarang bukan baris tabel
            if self.in_table and not re.match(r'^\s*\|', line):
                self._render_table()
                self.in_table = False

            # Cek fence code start
            if re.match(r'^```', line.strip()) and not self.in_code:
                self.handle_code_fence(line)
                i += 1
                continue

            # Heading
            if re.match(r'^#{1,6}\s+', line):
                self.handle_heading(line)
                i += 1
                continue

            # HR
            if line.strip() == '---':
                self.handle_hr(line)
                i += 1
                continue

            # Blockquote
            if line.startswith('>'):
                self.handle_blockquote(line)
                i += 1
                continue

            # Bullet
            if re.match(r'^[-*+]\s+', line):
                self.handle_bullet(line)
                i += 1
                continue

            # Ordered
            if re.match(r'^\d+\.\s+', line):
                self.handle_ordered(line)
                i += 1
                continue

            # Tabel
            if re.match(r'^\s*\|', line):
                self.handle_table(line)
                i += 1
                continue

            # Image Markdown: ![alt](path)
            m_img = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line.strip())
            if m_img:
                self.handle_image(m_img.group(1), m_img.group(2))
                i += 1
                continue

            # Kosong → spasi
            if not line.strip():
                # Jangan bikin paragraph kosong berlebihan
                if i + 1 < n and not lines[i+1].strip():
                    pass
                else:
                    self.doc.add_paragraph()
                i += 1
                continue

            # Paragraph biasa
            self.handle_paragraph(line)
            i += 1

        if self.in_table:
            self._render_table()
            self.in_table = False


# ─────────────────────────────────────────────────────────────────────────
# COVER & DAFTAR ISI
# ─────────────────────────────────────────────────────────────────────────

def add_cover(doc):
    # Spacer atas
    for _ in range(3):
        doc.add_paragraph()

    # Judul besar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('E-TICKETING')
    r.bold = True
    r.font.size = Pt(48)
    r.font.color.rgb = INK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('HELPDESK')
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = INK

    doc.add_paragraph()
    doc.add_paragraph()

    # Garis kuning tebal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('██████████████████████████████')
    r.font.color.rgb = YELLOW
    r.font.size = Pt(20)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LAPORAN UJIAN AKHIR SEMESTER (UAS)')
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Mata Kuliah Aplikasi Mobile — Semester 4')
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Jenis: UAS Teori')
    r.italic = True
    r.font.size = Pt(12)

    # Spacer besar
    for _ in range(6):
        doc.add_paragraph()

    # Info identitas (tabel kecil)
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    for row, (k, v) in enumerate([
        ('Nama Aplikasi',  'E-Ticketing Helpdesk'),
        ('Platform',       'Flutter 3.x (Android & iOS)'),
        ('Backend',        'Supabase (PostgreSQL + Auth + Storage)'),
        ('Arsitektur',     'MVVM + Provider + GoRouter'),
    ]):
        c0, c1 = table.rows[row].cells
        c0.width = Cm(4)
        c1.width = Cm(12)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)

    # Page break
    doc.add_page_break()


def add_toc(doc):
    """Daftar isi manual (TOC otomatis sulit saat build offline)."""
    p = doc.add_paragraph()
    r = p.add_run('DAFTAR ISI')
    r.bold = True
    r.font.size = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    entries = [
        ('Kata Pengantar',                                    'i'),
        ('BAB 1  Pendahuluan',                                '1'),
        ('  1.1  Latar Belakang',                             '1'),
        ('  1.2  Tujuan Pembuatan',                           '2'),
        ('  1.3  Ruang Lingkup',                              '2'),
        ('  1.4  User Persona',                               '3'),
        ('  1.5  Metodologi',                                 '4'),
        ('  1.6  Sistematika Penulisan',                      '4'),
        ('BAB 2  Arsitektur Sistem',                          '5'),
        ('  2.1  Arsitektur High-Level',                      '5'),
        ('  2.2  Tech Stack Lengkap',                         '6'),
        ('  2.3  Pola MVVM (Model–View–ViewModel)',          '7'),
        ('  2.4  Flow Diagram — Login',                       '8'),
        ('  2.5  Flow Diagram — Buat Tiket',                  '9'),
        ('  2.6  Flow Diagram — Ubah Status Tiket',          '10'),
        ('  2.7  Flow Diagram — Notifikasi Realtime',       '11'),
        ('  2.8  Struktur Folder Proyek',                    '12'),
        ('  2.9  Routing (GoRouter)',                        '13'),
        ('  2.10 Lifecycle Aplikasi',                        '14'),
        ('  2.11 Use Case Diagram',                         '15'),
        ('  2.12 Activity Diagram — Buat Tiket → Selesai',  '16'),
        ('  2.13 Component Diagram',                        '17'),
        ('BAB 3  Database',                                  '18'),
        ('  3.1  Platform',                                 '18'),
        ('  3.2  Daftar Tabel',                             '18'),
        ('  3.3  Skema Lengkap',                            '19'),
        ('  3.4  ERD (Entity Relationship Diagram)',        '20'),
        ('  3.5  Relasi Kunci (Foreign Key Detail)',        '21'),
        ('  3.6  Data Riil (Hasil Seed)',                   '21'),
        ('  3.7  Contoh Query',                             '22'),
        ('  3.8  Indexes',                                  '23'),
        ('  3.9  Row Level Security (RLS)',                 '23'),
        ('  3.10 State Machine Lifecycle',                  '24'),
        ('  3.11 Storage Bucket',                           '25'),
        ('BAB 4  Backend API',                               '26'),
        ('  4.1  Arsitektur REST',                           '26'),
        ('  4.2  Endpoint per Fitur',                        '26'),
        ('  4.3  SupabaseService',                           '27'),
        ('  4.4  Error Handling',                            '28'),
        ('BAB 5  UI / UX Aplikasi',                          '29'),
        ('  5.1  Design Language',                           '29'),
        ('  5.2  Daftar Layar',                              '29'),
        ('  5.3  Wireframe',                                 '30'),
        ('  5.4  Komponen UI',                               '32'),
        ('  5.5  Bottom Navigation',                         '32'),
        ('  5.6  Dark Mode',                                 '32'),
        ('  5.7  Accessibility',                             '32'),
        ('BAB 6  Fitur Aplikasi',                            '33'),
        ('  6.1  Role & Hak Akses',                          '33'),
        ('  6.2  Fitur User',                                '33'),
        ('  6.3  Fitur Helpdesk',                            '34'),
        ('  6.4  Fitur Admin',                               '34'),
        ('  6.5  Matriks Aksi',                              '35'),
        ('  6.6  Skenario Penggunaan',                       '35'),
        ('  6.7  Statistik Dashboard',                       '36'),
        ('BAB 7  Video Tutorial',                            '37'),
        ('  7.1  Spesifikasi',                               '37'),
        ('  7.2  Tools',                                     '37'),
        ('  7.3  Persiapan',                                 '37'),
        ('  7.4  Script Scene-by-Scene',                     '38'),
        ('  7.5  Checklist',                                 '40'),
        ('  7.6  Tips',                                      '40'),
        ('BAB 8  Panduan Deploy & Instalasi',                '41'),
        ('  8.1  Prasyarat',                                 '41'),
        ('  8.2  Clone',                                     '41'),
        ('  8.3  Install Dependencies',                      '41'),
        ('  8.4  Setup Supabase',                            '41'),
        ('  8.5  Setup Storage',                             '42'),
        ('  8.6  Policies',                                  '42'),
        ('  8.7  Jalankan Aplikasi',                         '42'),
        ('  8.8  Login Demo',                                '43'),
        ('  8.9  Build APK',                                 '43'),
        ('  8.10 Troubleshooting',                          '43'),
        ('  8.11 Reset Database',                            '44'),
        ('BAB 9  Dokumentasi Kode',                          '45'),
        ('  9.1  Statistik',                                 '45'),
        ('  9.2  Modul Kunci',                               '45'),
        ('  9.3  Models',                                    '47'),
        ('  9.4  Snippet',                                   '48'),
        ('  9.5  Alasan Teknis',                             '49'),
        ('  9.6  Testing',                                   '49'),
        ('  9.7  Limitasi',                                  '50'),
        ('  9.8  File Penting',                              '50'),
        ('Daftar Pustaka',                                   '51'),
        ('Lampiran',                                         '52'),
    ]

    for title, page in entries:
        p = doc.add_paragraph()
        # Tab leader ke kanan
        from docx.shared import Cm
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT, 2)  # 2 = dot leader
        r = p.add_run(title)
        r.font.size = Pt(11)
        r.add_tab()
        r2 = p.add_run(page)
        r2.font.size = Pt(11)

    doc.add_page_break()


def add_kata_pengantar(doc):
    p = doc.add_paragraph()
    r = p.add_run('KATA PENGANTAR')
    r.bold = True
    r.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    paragraphs = [
        "Segala puji dan syukur penulis panjatkan ke hadirat Tuhan Yang Maha Esa, "
        "atas rahmat dan karunia-Nya sehingga penulis dapat menyelesaikan laporan "
        "Ujian Akhir Semester (UAS) Teori mata kuliah Aplikasi Mobile ini dengan "
        "baik. Laporan ini mendokumentasikan aplikasi E-Ticketing Helpdesk, sebuah "
        "sistem tiket bantuan internal kantor berbasis Flutter dan Supabase.",

        "Penulis menyadari bahwa laporan ini tidak akan tersusun tanpa bantuan dari "
        "berbagai pihak. Oleh karena itu, penulis mengucapkan terima kasih kepada "
        "dosen pengampu, keluarga, dan teman-teman yang telah memberikan dukungan "
        "dan semangat selama pengerjaan tugas ini.",

        "Penulis menyadari sepenuhnya bahwa laporan ini masih jauh dari kata sempurna. "
        "Oleh karena itu, penulis dengan tangan terbuka menerima kritik dan saran "
        "yang membangun dari para pembaca demi perbaikan di masa datang.",

        "Akhir kata, semoga laporan ini dapat bermanfaat bagi siapa pun yang membaca, "
        "khususnya bagi yang sedang belajar pengembangan aplikasi mobile.",
    ]

    for txt in paragraphs:
        p = doc.add_paragraph(txt)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('Penulis')
    r.italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run('GordicProject')
    r2.bold = True

    doc.add_page_break()


def add_daftar_pustaka(doc):
    p = doc.add_paragraph()
    r = p.add_run('DAFTAR PUSTAKA')
    r.bold = True
    r.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    refs = [
        'Flutter Team. (2024). *Flutter Documentation*. https://docs.flutter.dev',
        'Google. (2024). *Dart Language Tour*. https://dart.dev/guides/language/language-tour',
        'Provider Authors. (2024). *Provider Package Documentation*. https://pub.dev/packages/provider',
        'Flutter Community. (2024). *go_router Documentation*. https://pub.dev/packages/go_router',
        'Supabase Inc. (2024). *Supabase Documentation*. https://supabase.com/docs',
        'Supabase Inc. (2024). *PostgREST API Reference*. https://postgrest.org',
        'Material Design. (2024). *Material Design 3*. https://m3.material.io',
        'Neo-Brutalism Web Design. (2024). *Design Resources*. https://brutalistwebsites.com',
        'Riverpod Authors. (2024). *State Management Patterns*. https://riverpod.dev',
        'PostgreSQL Global Development Group. (2024). *PostgreSQL 15 Documentation*. https://www.postgresql.org/docs/15/',
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent   = Cm(0.75)
        p.paragraph_format.space_after   = Pt(8)
        p.paragraph_format.line_spacing  = 1.25
        r = p.add_run(f'{i}.  ')
        r.bold = True
        # Parse simple markdown *..* jadi italic
        parts = re.split(r'\*([^*]+)\*', ref)
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            if idx % 2 == 1:
                run.italic = True

    doc.add_page_break()


def add_lampiran(doc):
    p = doc.add_paragraph()
    r = p.add_run('LAMPIRAN')
    r.bold = True
    r.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(
        "Pada bagian lampiran ini, dilampirkan file-file yang mendukung laporan, antara lain: "
        "source code aplikasi, skema database SQL, dan dokumen pendukung lainnya. Seluruh file "
        "tersebut tersedia di repository GitHub penulis dan folder `laporan/` dari proyek."
    ).italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Lampiran A — Struktur Folder Proyek')
    r.bold = True
    r.font.size = Pt(13)

    tree = """eTicketing/
├── lib/                          # Source code Flutter
│   ├── main.dart
│   ├── core/                     # Konfigurasi & utility
│   ├── data/                     # Models + dummy
│   ├── presentation/             # Screens, widgets, providers
│   └── services/                 # Supabase service
├── test/                         # Unit/widget tests
├── database_setup.sql            # Schema + seed
├── pubspec.yaml                  # Dependencies
├── android/                      # Android config
├── ios/                          # iOS config
└── laporan/                      # 📁 Laporan UAS
    ├── README.md
    ├── 01-pendahuluan.md
    ├── 02-arsitektur.md
    ├── 03-database.md
    ├── 04-api.md
    ├── 05-uiux.md
    ├── 06-fitur.md
    ├── 07-video-tutorial.md
    ├── 08-panduan-deploy.md
    └── 09-dokumentasi-kode.md"""

    # Render as code block
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, CODE_BG)
    cell.text = ''
    for i, line in enumerate(tree.split('\n')):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(9)


# ─────────────────────────────────────────────────────────────────────────
# HEADER & FOOTER
# ─────────────────────────────────────────────────────────────────────────
def setup_header_footer(doc):
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # Header
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('E-Ticketing Helpdesk — Laporan UAS Aplikasi Mobile')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.italic = True

    # Footer dengan nomor halaman
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # field kode untuk nomor halaman
    def add_page_field(par):
        run = par.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)

    r = p.add_run('— ')
    r.font.size = Pt(9)
    add_page_field(p)
    r = p.add_run(' —')
    r.font.size = Pt(9)


# ─────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    setup_header_footer(doc)
    add_cover(doc)
    add_kata_pengantar(doc)
    add_toc(doc)

    # Tiap bab — page break dulu, lalu judul bab, lalu isi
    for filename, bab_title in FILES:
        path = BASE / filename
        if not path.exists():
            print(f"[WARN]  Missing: {filename}")
            continue
        text = path.read_text(encoding='utf-8')

        # Strip leading H1 kalau sama dengan bab_title (heading sudah di cover-bab)
        lines = text.split('\n')

        # Page break + judul bab
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(bab_title)
        r.bold = True
        r.font.size = Pt(22)
        r.font.color.rgb = INK
        # Garis bawah
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run('═' * 40)
        r2.font.color.rgb = YELLOW
        r2.font.size = Pt(10)

        doc.add_paragraph()

        # Strip leading H1 dari markdown (heading sudah dibuat di atas)
        content_lines = []
        skipped_first_h1 = False
        for line in lines:
            if not skipped_first_h1 and re.match(r'^#\s+', line):
                skipped_first_h1 = True
                continue  # skip
            content_lines.append(line)

        # Parse
        parser = MdParser(doc)
        parser.feed(content_lines)

    # Penutup: daftar pustaka + lampiran
    add_daftar_pustaka(doc)
    add_lampiran(doc)

    # Coba save ke OUT; kalau file sedang dikunci (terbuka di Word),
    # fallback ke OUT_FALLBACK supaya build tidak gagal total.
    try:
        doc.save(str(OUT))
        print(f"[OK]  Dokumen tersimpan di:")
        print(f"   {OUT}")
        print(f"   Ukuran: {OUT.stat().st_size:,} bytes")
    except PermissionError:
        doc.save(str(OUT_FALLBACK))
        print(f"[WARN] File utama sedang dikunci — simpan ke fallback:")
        print(f"   {OUT_FALLBACK}")
        print(f"   Ukuran: {OUT_FALLBACK.stat().st_size:,} bytes")


if __name__ == '__main__':
    main()
